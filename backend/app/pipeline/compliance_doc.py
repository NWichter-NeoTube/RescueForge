"""DIN 14095 / FKS Compliance Documentation Generator.

Generates a .docx compliance report listing which DIN standards the plan
adheres to, which elements were generated, and which items require manual
input (cannot be derived from the DXF upload).
"""

import logging
from pathlib import Path

from app.models.schemas import FloorPlanData, RoomPolygon
from app.utils.translations import Locale, t

logger = logging.getLogger(__name__)


def generate_compliance_doc(
    floor_plan: FloorPlanData,
    rooms: list[RoomPolygon],
    output_path: Path,
    language: str = "en",
) -> Path:
    """Generate a DIN 14095 compliance report as .docx.

    Args:
        floor_plan: Parsed floor plan data with detected features.
        rooms: Classified room polygons.
        output_path: Where to save the .docx file.
        language: "en" or "de".

    Returns:
        Path to the generated .docx file.
    """
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.shared import Inches, Pt, RGBColor

    lang: Locale = "de" if language == "de" else "en"

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(10)

    # ── Title ──────────────────────────────────────────────────
    title = doc.add_heading(t("compliance.title", lang), level=0)
    title.runs[0].font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph(f"{t('title.object', lang)}: {floor_plan.filename}")
    doc.add_paragraph(f"{t('title.floor', lang)}: {floor_plan.floor_label or 'EG'}")
    doc.add_paragraph("")

    # ── Standards Referenced ───────────────────────────────────
    doc.add_heading(t("compliance.standards", lang), level=1)
    standards = [
        "DIN 14095:2007-05 — Feuerwehrpläne für bauliche Anlagen",
        "DIN 14034-6:2024-06 — Graphische Symbole für das Feuerwehrwesen (Teil 6)",
        "FKS-Richtlinie — Orientierungspläne (Schweiz)",
    ]
    for s in standards:
        doc.add_paragraph(s, style="List Bullet")

    doc.add_paragraph("")

    # ── Generated Elements Table ──────────────────────────────
    doc.add_heading(t("compliance.elements", lang), level=1)

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    # Header row
    hdr = table.rows[0].cells
    hdr[0].text = t("compliance.element", lang)
    hdr[1].text = t("compliance.din_reference", lang)
    hdr[2].text = t("compliance.status", lang)
    for cell in hdr:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)

    # Element status entries
    elements = _build_element_status(floor_plan, rooms, lang)
    for elem_name, din_ref, status in elements:
        row = table.add_row().cells
        row[0].text = elem_name
        row[1].text = din_ref
        row[2].text = status
        # Color status
        for paragraph in row[2].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)
                if status == t("compliance.implemented", lang):
                    run.font.color.rgb = RGBColor(0, 128, 0)
                elif status == t("compliance.partial", lang):
                    run.font.color.rgb = RGBColor(255, 165, 0)
                else:
                    run.font.color.rgb = RGBColor(200, 0, 0)

    doc.add_paragraph("")

    # ── Manual Input Required ─────────────────────────────────
    doc.add_heading(t("compliance.manual_input", lang), level=1)
    doc.add_paragraph(
        t("compliance.description", lang) + ":"
        if lang == "de"
        else "The following items are required by DIN 14095 but cannot be "
             "automatically derived from the DXF floor plan data. They must "
             "be added manually by the plan creator:"
    )

    manual_items = _build_manual_items(lang)
    manual_table = doc.add_table(rows=1, cols=2)
    manual_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    manual_table.style = "Table Grid"

    mhdr = manual_table.rows[0].cells
    mhdr[0].text = t("compliance.element", lang)
    mhdr[1].text = t("compliance.reason", lang)
    for cell in mhdr:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)

    for item_name, reason in manual_items:
        row = manual_table.add_row().cells
        row[0].text = item_name
        row[1].text = reason

    # Save
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    logger.info("Compliance document generated: %s", output_path)
    return output_path


def _build_element_status(
    floor_plan: FloorPlanData,
    rooms: list[RoomPolygon],
    lang: Locale,
) -> list[tuple[str, str, str]]:
    """Build the list of (element, DIN reference, status) for the compliance table."""
    implemented = t("compliance.implemented", lang)
    partial = t("compliance.partial", lang)
    not_available = t("compliance.not_available", lang)

    has_rooms = len(rooms) > 0
    has_walls = len(floor_plan.walls) > 0
    has_doors = len(floor_plan.doors) > 0
    has_stairs = len(floor_plan.stairs) > 0
    has_fire_walls = len(floor_plan.fire_walls) > 0
    has_fire_doors = len(floor_plan.fire_doors) > 0
    has_windows = len(floor_plan.windows) > 0
    has_sprinkler = floor_plan.has_sprinkler

    elements = [
        # (Element name, DIN reference, status)
        (t("title.orientation_plan", lang), "DIN 14095 §4", implemented if has_rooms else partial),
        (t("title.cover_sheet", lang), "DIN 14095 §4.1", implemented),
        (t("title.situation_plan", lang), "DIN 14095 §4.2", implemented),
        (t("color.walls", lang) + " (RAL 9004)", "DIN 14095 §5.1", implemented if has_walls else not_available),
        (t("legend.fire_wall", lang) + " (RAL 3000)", "DIN 14095 §5.2", implemented if has_fire_walls else not_available),
        (t("legend.fire_door", lang), "DIN 14095 §5.3", implemented if has_fire_doors else not_available),
        (t("color.doors", lang), "DIN 14095 §5.1", implemented if has_doors else not_available),
        (t("legend.vertical_escape", lang), "DIN 14095 §6.1", implemented if has_stairs else partial),
        (t("legend.horizontal_escape", lang), "DIN 14095 §6.2", implemented if has_rooms else partial),
        (t("symbol.north_arrow", lang), "DIN 14095 §7.1", implemented),
        (t("symbol.scale_bar", lang), "DIN 14095 §7.2", implemented),
        (t("symbol.smoke_detector", lang), "DIN 14034-6", implemented if has_rooms else not_available),
        (t("symbol.manual_call_point", lang), "DIN 14034-6", implemented),
        (t("symbol.bma", lang), "DIN 14034-6", implemented),
        (t("symbol.fire_access", lang), "DIN 14034-6", implemented),
        (t("symbol.key_depot", lang), "DIN 14034-6", implemented),
        (t("symbol.gas_shutoff", lang), "DIN 14034-6", implemented),
        (t("symbol.water_shutoff", lang), "DIN 14034-6", implemented),
        (t("symbol.electricity_shutoff", lang), "DIN 14034-6", implemented),
        (t("symbol.stair_direction", lang), "DIN 14034-6", implemented if has_stairs else not_available),
        (t("symbol.elevator", lang), "DIN 14034-6", implemented if any(r.room_type.value == "elevator" for r in rooms) else not_available),
        (t("symbol.wall_hydrant", lang), "DIN 14034-6", implemented if has_sprinkler else not_available),
        (t("symbol.sprinkler", lang), "DIN 14034-6", implemented if has_sprinkler else not_available),
        (t("symbol.rwa", lang), "DIN 14034-6", partial),
        (t("symbol.assembly_point", lang), "DIN 14034-6", implemented),
        (t("symbol.erstinfo", lang), "DIN 14034-6:2024", implemented),
        ("10m " + t("symbol.scale_bar", lang) + " Grid", "DIN 14095 §5.3", implemented),
        (t("legend.non_passable", lang), "DIN 14095 §5.4", implemented if any(r.room_type.value in ("technical", "server_room", "garage") for r in rooms) else not_available),
    ]

    # Windows
    if has_windows:
        elements.append(("Windows / Fenster", "DIN 14095 §5.1", implemented))

    return elements


def _build_manual_items(lang: Locale) -> list[tuple[str, str]]:
    """Build the list of items requiring manual input."""
    if lang == "de":
        return [
            ("Gefahrstoff-Standorte und Mengen",
             "Gefahrstoffe sind im DXF nicht kodiert; müssen vom Betreiber erfasst werden"),
            ("Gas-/Wasser-/Elektro-Absperrstellen (exakte Positionen)",
             "Nur symbolisch platziert; genaue Positionen erfordern Begehung"),
            ("Sprinkler-Leitungsverläufe und Ventilpositionen",
             "Sprinkler-Erkennung ist binär; Leitungsdetails nicht im DXF"),
            ("Brandwiderstandsklasse der Wände (T30/T60/T90/F30/F60/F90)",
             "Erfordert Bauunterlagen; nicht aus Geometrie ableitbar"),
            ("Kontaktdaten Gebäudeeigentümer",
             "Personenbezogene Daten nicht im CAD-Plan"),
            ("Taktische Hinweise der Feuerwehr",
             "Erfordert Abstimmung mit zuständiger Feuerwehr"),
            ("RWA-Bedienstellen (exakte Positionen)",
             "Nur Stairwell-basiert platziert; exakte Positionen erfordern Dokumentation"),
            ("Schlüsseldepot-Zugangscode",
             "Sicherheitsrelevante Information; nicht im DXF"),
            ("Sammelplatz-Standort (falls nicht im DXF)",
             "Generisch im Situationsplan platziert; erfordert Bestätigung"),
            ("PV-Anlage Details (Leistung, Wechselrichter-Standort)",
             "Photovoltaik-Erkennung ist binär; technische Details nicht im DXF"),
            ("Objektfunkanlage",
             "Funkanlagen-Information nicht im CAD-Plan kodiert"),
            ("Notfall-Kontaktliste",
             "Personenbezogene Daten; müssen vom Betreiber bereitgestellt werden"),
        ]
    else:
        return [
            ("Hazardous material locations & quantities",
             "Hazardous materials are not encoded in DXF; must be surveyed on-site"),
            ("Gas/water/electricity shutoff positions (exact)",
             "Placed symbolically; exact positions require on-site inspection"),
            ("Sprinkler pipe routes & valve locations",
             "Sprinkler detection is binary; pipe routing details not in DXF"),
            ("Fire resistance class of walls (T30/T60/T90/F30/F60/F90)",
             "Requires building documentation; not derivable from geometry"),
            ("Building owner contact information",
             "Personal data not encoded in CAD plans"),
            ("Fire department tactical notes",
             "Requires coordination with responsible fire department"),
            ("RWA control panel positions (exact)",
             "Placed near stairwells; exact positions require documentation"),
            ("Key safe code / access details",
             "Security-sensitive information; not in DXF"),
            ("Assembly point location (if not in DXF)",
             "Placed generically in situation plan; requires confirmation"),
            ("PV system details (capacity, inverter location)",
             "PV detection is binary; technical details not in DXF"),
            ("Object radio system (Objektfunkanlage)",
             "Radio system information not encoded in CAD plans"),
            ("Emergency contact list",
             "Personal data; must be provided by building operator"),
        ]
